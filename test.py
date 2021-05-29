import os
import shutil
import docx
from docx.shared import RGBColor, Pt

application_file='C:\\Users\\Nico\\Desktop\\py练习\\结关申请.docx'


open_application_file=docx.Document(application_file)

for root,dirs,files in os.walk('C:\\Users\\Nico\\Desktop\\out'):
    for dir in dirs:
        print('dir: ',dir)
        for p in open_application_file.paragraphs:
            p.text = p.text.replace('123456789', dir[4:6])
            p.text = p.text.replace('yyyy-mm-dd', '2021-05-24')
        for paragraph in open_application_file.paragraphs:
            for run in paragraph.runs:
                run.font.name = '宋体'
                run.font.size = Pt(14)
        open_application_file.save('C:\\Users\\Nico\\Desktop\\out\\'+dir+'\\'+'结关申请.docx')