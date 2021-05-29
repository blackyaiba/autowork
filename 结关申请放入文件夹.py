import os
import docx
from docx.shared import RGBColor, Pt

application_file = '\\\\192.168.128.27\\DFQD-Shared\\市场服务部\\物流组\\关务\\（日常）强行结关（目前每月下旬提交之前数据）\\模板\\结关申请模板.docx'


for root, dirs, files in os.walk('C:\\Users\\shilanqian\\Desktop\\5月结关'):
    for dir in dirs:
        open_application_file = docx.Document(application_file)
        for p in open_application_file.paragraphs:
            p.text = p.text.replace('123456789', str(dir)[4:13])
            p.text = p.text.replace('yyyy-mm-dd', '2021-05-29')
        for paragraph in open_application_file.paragraphs:
            for run in paragraph.runs:
                run.font.name = '宋体'
                run.font.size = Pt(14)
        open_application_file.save('C:\\Users\\shilanqian\\Desktop\\5月结关\\' + dir + '\\' + '结关申请.docx')

